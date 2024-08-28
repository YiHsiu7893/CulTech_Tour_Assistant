import requests
from bs4 import BeautifulSoup 
from docx import Document


def crawl_content(url):
    #取出網頁原始碼
    resp = requests.get(url)
    soup = BeautifulSoup(resp.text,"html.parser") 
    #print(soup)

    # 取出標題
    h = soup.find_all("h1")
    if h:
        header = h[0].get_text(strip=True) + "\n"
    else:
        header = None

    # 取出文字內容
    content = ""
    paragraphs = soup.find_all("p")
    for p in paragraphs:
        ptext = p.get_text(strip=True)
        if ptext != "©2022 Peiguihall":
            content += ptext + "。\n"

    return header, content


if __name__ == '__main__':
    doc = Document()
    
    # 所有培桂堂相關網頁
    urls = [
    "https://www.peiguihall.org.tw/tw/about/page/%E9%97%9C%E6%96%BC%E5%9F%B9%E6%A1%82%E5%A0%82",
    "https://www.peiguihall.org.tw/tw/about/page/%E5%BB%BA%E7%AF%89%E4%BB%8B%E7%B4%B9",
    "https://www.peiguihall.org.tw/tw/about/page/%E5%BA%AD%E9%99%A2%E8%88%87%E4%BC%91%E6%86%A9",
    "https://www.peiguihall.org.tw/tw/visit/page/%E5%8F%83%E8%A7%80%E8%B3%87%E8%A8%8A",
    "https://www.peiguihall.org.tw/tw/visit/page/%E5%B0%8E%E8%A6%BD%E6%9C%8D%E5%8B%99",
    "https://www.peiguihall.org.tw/tw/visit/page/%E4%BA%A4%E9%80%9A%E8%B3%87%E8%A8%8A",
    "https://www.peiguihall.org.tw/tw/viewpoint"]


    for u in urls:
        header, content = crawl_content(u)

        # 加標題
        if header:
            doc.add_heading(header, level=2)
        # 加內文
        doc.add_paragraph(content)
        doc.add_page_break()

    doc.save("peigui_hall.docx")